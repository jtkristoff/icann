#!/usr/bin/env python3

import argparse
import webvtt

from docx import Document
from docx.shared import Inches
from docx.shared import Pt

parser = argparse.ArgumentParser()
parser.add_argument( '-f', default='web.vtt' )  # WebVTT file to parse
parser.add_argument( '-l', default=30 )         # default name string length
parser.add_argument( '-o', default='web.docx' ) # output file
parser.add_argument( '-t', default='Title' )
args = parser.parse_args()

document = Document()
document.add_heading(args.t, 0)

speaker = ''
previous_speaker = ''
speech = ''

for caption in webvtt.read(args.f):
    try:
        # do we have text and a speaker?
        dialog = caption.text.split(':')[1].strip()
        speaker = caption.text.split(':')[0].strip()

        # remove any appended org appended after a dash
        speaker = speaker.split('-')[0].strip()

        # speaker's name string length to protect formatting
        speaker = speaker[:args.l]
    except:
        # if we just have text, the speaker hasn't changed
        speaker = previous_speaker
        dialog = caption.text.split(':')[0].strip()

    # if this the same speaker, combine this dialog into a larger speech
    if speaker == previous_speaker:
        if speech:
            speech = speech + ' ' + dialog
        else:
            speech = dialog
        continue
    else:
         # new speaker, so output the last speech
         p = document.add_paragraph()
         p.paragraph_format.first_line_indent = Inches(-0.25)
         p.add_run(previous_speaker.upper() + '  ').bold = True
         p.add_run(speech)
         speech = dialog

    previous_speaker = speaker

# output last speaker+speech
p = document.add_paragraph()
p.paragraph_format.first_line_indent = Inches(-0.25)
p.add_run(previous_speaker.upper() + '  ').bold = True
p.add_run(speech)

document.save(args.o)
